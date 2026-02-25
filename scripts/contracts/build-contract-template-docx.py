#!/usr/bin/env python3
"""
Build a styled DOCX contract template from the text source.

Usage:
  python3 scripts/contracts/build-contract-template-docx.py
"""

from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "config/contracts/templates/contrato_arrendamiento_template_v1.txt"
OUT = ROOT / "config/contracts/templates/contrato_arrendamiento_template_v1.docx"
LOGO = ROOT / "config/contracts/assets/uaf-header.png"

PRICE_ROWS = [
    ("Aseo", "Aseo General Básico", "1,145 UF + IVA", "c/u"),
    ("Aseo", "Aseo General Intenso", "1,527 UF + IVA", "c/u"),
    ("Pintura", "Pintura Guardapolvos", "0,046 UF + IVA", "ml"),
    ("Pintura", "Pintura Muros", "0,573 UF + IVA", "c/u"),
    ("Pintura", "Pintura Interior de Closet Chico", "0,382 UF + IVA", "c/u"),
    ("Pintura", "Pintura Interior de Closet Grande", "0,573 UF + IVA", "c/u"),
    ("Pintura", "Pintura Puerta", "0,382 UF + IVA", "c/u"),
    ("Pintura", "Pintura Marco de Puerta", "0,305 UF + IVA", "c/u"),
    ("Pintura", "Pintura Cielo", "1,145 UF + IVA", "por habitación"),
    ("Gasfitería", "Cambio de Flexible", "0,305 UF + IVA", "c/u"),
    ("Gasfitería", "Cambio de Flexible Ducha", "0,229 UF + IVA", "c/u"),
    ("Gasfitería", "Cambio de Challa Ducha", "0,191 UF + IVA", "c/u"),
    ("Gasfitería", "Destape Desagües y Sifón", "0,153 UF + IVA", "c/u"),
    ("Gasfitería", "Cambio de Sifón", "0,573 UF + IVA", "c/u"),
    ("Gasfitería", "Cambio Tapa WC", "0,954 UF + IVA", "c/u"),
    ("Gasfitería", "Flaper WC", "0,763 UF + IVA", "c/u"),
    ("Gasfitería", "Monomando Cocina", "1,145 UF + IVA", "c/u"),
    ("Gasfitería", "Monomando Baño", "0,763 UF + IVA", "c/u"),
    ("Cerrajería", "Cambio de Chapa Acceso Principal", "0,763 UF + IVA", "c/u"),
    ("Carpintería", "Cambio de Puerta Interior", "1,718 UF + IVA", "c/u"),
    ("Rep. Generales", "Reparación forado en tabique (r<6cm)", "0,954 UF + IVA", "c/u"),
    ("Rep. Generales", "Cambio de piso", "0,704 UF + IVA", "m2"),
    ("Rep. Generales", "Pérdida o daño tarjeta de acceso", "2 UF + IVA", "unidad"),
    ("Rep. Generales", "Pérdida o daño llave acceso y/o bodega", "0,185 UF + IVA", "unidad"),
]

FURNISHED_PRICE_ROWS = [
    ("Mobiliario", "Cama 2 Plazas", "8,25 UF + IVA", "unidad"),
    ("Mobiliario", "Velador pieza principal", "2 UF + IVA", "c/u"),
    ("Mobiliario", "Respaldo cama pieza principal", "2 UF + IVA", "unidad"),
    ("Mobiliario", "Cubrecolchón cama 2 plazas", "0,75 UF + IVA", "unidad"),
    ("Mobiliario", "Lámpara velador", "1 UF + IVA", "c/u"),
    ("Mobiliario", "Mueble/arrimo", "1,5 UF + IVA", "unidad"),
    ("Mobiliario", "Camarote pieza chica", "5 UF + IVA", "unidad"),
    ("Mobiliario", "Colchón 1 plaza", "3 UF + IVA", "c/u"),
    ("Mobiliario", "Cubrecolchón 1 Plaza", "0,5 UF + IVA", "c/u"),
    ("Mobiliario", "Estante", "1,5 UF + IVA", "unidad"),
    ("Mobiliario", "Velador pieza chica", "1 UF + IVA", "unidad"),
    ("Mobiliario", "Reposición Microondas", "1,6286 UF + IVA", "c/u"),
    ("Mobiliario", "Reposición Juego de Terraza", "1,9227 UF + IVA", "c/u"),
    ("Mobiliario", "Reposición Pisos de Cocina", "1,242 UF + IVA", "c/u"),
]


def configure_section(section) -> None:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.0)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "666666")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_uaf_banner(doc: Document, page_break_before: bool = False) -> None:
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_paragraph.paragraph_format.page_break_before = page_break_before
    logo_paragraph.paragraph_format.space_before = Pt(0)
    logo_paragraph.paragraph_format.space_after = Pt(4)
    if LOGO.exists():
        logo_paragraph.add_run().add_picture(str(LOGO), width=Cm(6.4))
    else:
        fallback = logo_paragraph.add_run("UNIDAD DE ANÁLISIS FINANCIERO")
        fallback.bold = True
        fallback.font.size = Pt(12)

    line_paragraph = doc.add_paragraph("")
    line_paragraph.paragraph_format.space_before = Pt(0)
    line_paragraph.paragraph_format.space_after = Pt(10)
    set_bottom_border(line_paragraph)


def normalize_spanish_text(text: str) -> str:
    replacements = {
        "rol unico": "rol único",
        "segun": "según",
        "numero": "número",
        "cedula": "cédula",
        "anos": "años",
        "ano": "año",
        "canerias": "cañerías",
        "danos": "daños",
        "Desagues": "Desagües",
        "Sifon": "Sifón",
    }
    out = text
    for old, new in replacements.items():
        out = out.replace(old, new)
    return out


def add_price_table(doc: Document, rows: list[tuple[str, str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ("ITEM", "DESCRIPCIÓN", "PRECIO UNITARIO ($)", "UNIDAD")
    widths = (Cm(2.7), Cm(6.8), Cm(3.7), Cm(2.3))

    header_row = table.rows[0]
    set_repeat_header(header_row)
    for idx, cell in enumerate(header_row.cells):
        cell.width = widths[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx < 2 else WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(headers[idx])
        run.bold = True
        run.font.size = Pt(11)
        shade_cell(cell, "E6E6E6")

    for item, desc, price, unit in rows:
        row_cells = table.add_row().cells
        values = (item, desc, price, unit)
        for idx, value in enumerate(values):
            cell = row_cells[idx]
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx < 2 else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            run.font.size = Pt(10.5)


def main() -> None:
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_section(doc.sections[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    heading_re = re.compile(r"^(PRIMERO|SEGUNDO|TERCERO|CUARTO|QUINTO|SEXTO|SEPTIMO|OCTAVO|NOVENO|DECIMO(\s+\w+)?)")
    declaration_heading_re = re.compile(r"^DECLARACION DE ")
    centered_tokens = {
        "CONTRATO DE ARRENDAMIENTO",
        "CON",
        "[[ARRENDADORA.RAZON_SOCIAL]]",
        "[[ARRENDATARIO.NOMBRE]]",
    }
    account_prefixes = (
        "Titular:",
        "RUT:",
        "Banco:",
        "Tipo de cuenta:",
        "Numero de cuenta:",
        "Correo de pago:",
    )

    i = 0
    declaration_count = 0
    while i < len(lines):
        text = lines[i].rstrip("\n")
        normalized = normalize_spanish_text(text)

        if normalized.startswith("DECLARACION DE "):
            add_uaf_banner(doc, page_break_before=declaration_count > 0)
            declaration_count += 1

            p = doc.add_paragraph("")
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(normalized)
            run.bold = True
            run.font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if normalized == "Tabla referencial de precios unitarios de reparaciones (UF + IVA):":
            title = doc.add_paragraph()
            title.paragraph_format.space_before = Pt(8)
            title.paragraph_format.space_after = Pt(4)
            title.add_run("Tabla referencial de precios unitarios de reparaciones (UF + IVA):").bold = True
            add_price_table(doc, PRICE_ROWS)
            if i + 1 < len(lines) and "Aseo General Basico" in lines[i + 1]:
                i += 2
            else:
                i += 1
            continue

        if normalized.startswith("Los siguientes items aplican solo a departamentos amoblados"):
            title = doc.add_paragraph()
            title.paragraph_format.space_before = Pt(8)
            title.paragraph_format.space_after = Pt(4)
            title.add_run("Los siguientes ítems aplican solo a departamentos amoblados (UF + IVA):").bold = True
            add_price_table(doc, FURNISHED_PRICE_ROWS)
            i += 1
            while i < len(lines) and lines[i].lstrip().startswith("- "):
                i += 1
            continue

        paragraph = doc.add_paragraph("")
        if not normalized.strip():
            paragraph.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(6)

        if normalized in centered_tokens:
            run = paragraph.add_run(normalized)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.bold = True
            if normalized == "CONTRATO DE ARRENDAMIENTO":
                run.font.size = Pt(16)
                paragraph.paragraph_format.space_after = Pt(14)
            elif normalized == "CON":
                run.font.size = Pt(12)
                run.italic = True
            else:
                run.font.size = Pt(13)
            i += 1
            continue

        if heading_re.match(normalized):
            run = paragraph.add_run(normalized)
            run.bold = True
            run.font.size = Pt(12)
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if normalized.startswith("_________________________"):
            run = paragraph.add_run(normalized)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.bold = True
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        if declaration_heading_re.match(normalized):
            run = paragraph.add_run(normalized)
            run.bold = True
            run.font.size = Pt(12)
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if re.match(r"^[a-z]\)", normalized):
            paragraph.paragraph_format.left_indent = Cm(0.6)
            paragraph.paragraph_format.first_line_indent = Cm(-0.2)
        elif normalized.startswith("- "):
            paragraph.paragraph_format.left_indent = Cm(0.8)
            paragraph.paragraph_format.first_line_indent = Cm(-0.2)

        if ":" in normalized and normalized.startswith(account_prefixes):
            label, rest = normalized.split(":", 1)
            label_run = paragraph.add_run(f"{label}:")
            label_run.bold = True
            paragraph.add_run(rest)
        else:
            paragraph.add_run(normalized)
        i += 1

    if doc.paragraphs and not doc.paragraphs[0].text.strip():
        first = doc.paragraphs[0]._element
        first.getparent().remove(first)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    main()
